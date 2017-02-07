
# install via pip: python-docx
# https://python-docx.readthedocs.io/en/latest/
import docx
import re

def paragraph_replace(paragraphs, search, replace):
    searchre = re.compile(search)
    for paragraph in paragraphs:
        paragraph_text = paragraph.text
        if paragraph_text:
            if searchre.search(paragraph_text):
                paragraph.text = re.sub(search, replace, paragraph_text)
    return paragraph_text


doc = docx.Document("Template.docx")
paragraph_replace(doc.paragraphs,"Name der grundbesitzenden Gesellschaft","ICH AG")

# problems with text in textboxes
# workaround needed
paragraph_replace(doc.paragraphs,"Telefax","FAX")

doc.save("Template_Modified.docx")






# gets all the text from the pdf
def getAllText(filename):
    doc = docx.Document(filename)
    fullText = []
    paragraph_replace(doc.paragraphs,"Name der grundbesitzenden Gesellschaft","ICH AG")
    for para in doc.paragraphs:
        fullText.append(para.text)

    doc.save("Template_Modified.docx")
    return '\n'.join(fullText)
